#!/usr/bin/env python3
"""Extract text from DOCX file. Supports paragraph ranges and flags for structure."""

import sys
import argparse
from docx import Document


def extract(path: str, show_structure: bool = False, include_tables: bool = True,
            start_para: int = 0, end_para: int | None = None) -> None:
    doc = Document(path)

    # Compute the effective range
    total_paras = len(doc.paragraphs)
    if start_para < 0:
        start_para = 0
    if end_para is None or end_para > total_paras:
        end_para = total_paras

    para_idx = 0
    heading_stack = []  # Track current heading context

    for para in doc.paragraphs:
        if para_idx < start_para:
            para_idx += 1
            continue
        if para_idx >= end_para:
            break

        style = para.style.name if para.style else "Normal"
        text = para.text

        if not text.strip():
            if show_structure:
                print()
            para_idx += 1
            continue

        if show_structure:
            if style.lower().startswith("heading"):
                level_match = __import__('re').search(r'\d+', style)
                level = int(level_match.group()) if level_match else 1
                # Maintain heading stack
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text.strip()))
                prefix = "#" * min(level, 6)
                print(f"\n{prefix} {text.strip()}")
            elif style.lower().startswith("toc"):
                print(f"\n[TOC] {text.strip()}")
            elif "list" in style.lower():
                print(f"  • {text.strip()}")
            else:
                prefix = "  " if heading_stack else ""
                # Print paragraph style as annotation if not Normal
                if style != "Normal":
                    print(f"{prefix}[{style}] {text.strip()}")
                else:
                    print(f"{prefix}{text.strip()}")
        else:
            if text.strip():
                print(text.strip())

        para_idx += 1

    # --- Tables ---
    if include_tables and doc.tables:
        print("\n" + "=" * 60)
        print("=== TABLES ===")
        print("=" * 60)
        for t_idx, table in enumerate(doc.tables):
            print(f"\n--- Table {t_idx + 1} ({len(table.rows)} rows × {len(table.columns)} cols) ---")
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
                print("  " + "  |  ".join(cells))
                if r_idx == 0 and len(table.rows) > 1:
                    print("  " + "-" * 40)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from DOCX")
    parser.add_argument("path", help="Path to DOCX file")
    parser.add_argument("--structure", action="store_true",
                        help="Show heading hierarchy and styles")
    parser.add_argument("--no-tables", action="store_true",
                        help="Skip table extraction")
    parser.add_argument("--start", type=int, default=0,
                        help="Start paragraph index (0-indexed)")
    parser.add_argument("--end", type=int, default=None,
                        help="End paragraph index (exclusive)")
    args = parser.parse_args()

    extract(args.path, args.structure, not args.no_tables, args.start, args.end)
