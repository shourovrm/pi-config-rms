#!/usr/bin/env python3
"""DOCX metadata and structural analysis. Always run this first to triage a DOCX file."""

import sys
import json
import re
from docx import Document


def estimate_pages(char_count: int, chars_per_page: int = 3000) -> int:
    """Rough page count estimate based on character count."""
    return max(1, round(char_count / chars_per_page))


def analyze(path: str) -> dict:
    doc = Document(path)

    # --- Count paragraphs and analyze text ---
    paragraphs = []
    heading_count = 0
    total_text_length = 0
    math_chars_total = 0
    empty_paragraphs = 0

    for para in doc.paragraphs:
        text = para.text
        style = para.style.name if para.style else "Normal"
        is_heading = style.lower().startswith("heading") or style.lower().startswith("toc")
        text_len = len(text)

        if is_heading:
            heading_count += 1
        if text_len == 0:
            empty_paragraphs += 1

        total_text_length += text_len

        # Math character heuristic (Unicode math ranges)
        math_chars = sum(
            1 for c in text
            if ('\u2200' <= c <= '\u22FF')      # Mathematical Operators
            or ('\u2100' <= c <= '\u214F')       # Letterlike Symbols
            or ('\u2190' <= c <= '\u21FF')       # Arrows
            or ('\u27C0' <= c <= '\u27EF')       # Misc Mathematical Symbols-A
            or ('\u2980' <= c <= '\u29FF')       # Misc Mathematical Symbols-B
            or ('\u2A00' <= c <= '\u2AFF')       # Supplemental Mathematical Operators
            or ('\u0370' <= c <= '\u03FF')       # Greek
            or ('\u1D400' <= c <= '\u1D7FF')     # Mathematical Alphanumeric Symbols
            or c in '∫∑∏√∂∇∞≈≠≤≥±×÷∈∉⊂⊃∪∩∧∨¬∀∃∅'
        )
        math_chars_total += math_chars

    # --- Check for OMML (Office Math Markup Language) elements ---
    body_xml = doc.element.body.xml
    oml_count = len(re.findall(r'<m:oMath[>\s]', body_xml))
    # Also try alternate namespace prefix
    if oml_count == 0:
        oml_count = len(re.findall(r'<[^>]*:oMath[>\s]', body_xml))

    # --- Count tables ---
    table_count = len(doc.tables)
    table_rows_total = sum(len(table.rows) for table in doc.tables)

    # --- Count embedded images ---
    image_extensions = set()
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            # Extract file extension from target
            target = rel.target_ref if hasattr(rel, 'target_ref') else str(rel.target_ref) if hasattr(rel, 'target_ref') else ''
            ext = rel.target_ref.split('.')[-1].lower() if hasattr(rel, 'target_ref') and '.' in rel.target_ref else 'unknown'
            image_extensions.add(ext)
    image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)

    # --- Extract headings as pseudo-TOC ---
    headings = []
    for para in doc.paragraphs:
        if para.style and para.style.name.lower().startswith("heading"):
            level_match = re.search(r'\d+', para.style.name)
            level = int(level_match.group()) if level_match else 1
            headings.append({
                "level": level,
                "title": para.text.strip()[:120],  # Truncate long headings
                "style": para.style.name,
            })

    # --- Try to get core properties ---
    core_props = {}
    try:
        core_props = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
        }
    except Exception:
        pass

    estimated_pages = estimate_pages(total_text_length)
    meaningful_paras = len(doc.paragraphs) - empty_paragraphs

    result = {
        "file": path,
        "estimated_pages": estimated_pages,
        "paragraph_count": len(doc.paragraphs),
        "meaningful_paragraphs": meaningful_paras,
        "empty_paragraphs": empty_paragraphs,
        "heading_count": heading_count,
        "table_count": table_count,
        "table_rows_total": table_rows_total,
        "image_count": image_count,
        "image_types": sorted(image_extensions),
        "oml_math_count": oml_count,
        "total_text_length": total_text_length,
        "math_char_density": round(math_chars_total / max(total_text_length, 1), 4),
        "headings": headings,
        "core_properties": core_props,
    }

    return result


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: docx_info.py <path>", file=sys.stderr)
        sys.exit(1)

    result = analyze(sys.argv[1])
    print(json.dumps(result, indent=2))
