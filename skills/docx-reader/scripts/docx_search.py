#!/usr/bin/env python3
"""Search within DOCX text content by keyword or regex."""

import sys
import re
import argparse
from docx import Document


def search(path: str, query: str, context_lines: int = 3, literal: bool = False) -> None:
    doc = Document(path)

    if literal:
        pattern = re.compile(re.escape(query), re.IGNORECASE)
    else:
        try:
            pattern = re.compile(query, re.IGNORECASE)
        except re.error as e:
            print(f"Invalid regex: {e}", file=sys.stderr)
            sys.exit(1)

    total_matches = 0
    MAX_PREVIEW = 200

    # --- Search paragraphs ---
    para_list = list(doc.paragraphs)
    for para_idx, para in enumerate(para_list):
        text = para.text
        if not text.strip():
            continue

        lines = text.split("\n")
        for line_num, line in enumerate(lines):
            if pattern.search(line):
                total_matches += 1
                style = para.style.name if para.style else "Normal"
                print(f"\n=== Paragraph {para_idx}, line {line_num + 1} [{style}] ===")

                # Show context lines within the SAME paragraph only
                start = max(0, line_num - context_lines)
                end = min(len(lines), line_num + context_lines + 1)
                for j in range(start, end):
                    marker = ">>>" if j == line_num else "   "
                    print(f"{marker} {lines[j][:MAX_PREVIEW]}")

    # --- Search tables ---
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                if pattern.search(text):
                    total_matches += 1
                    print(f"\n=== Table {t_idx + 1}, row {r_idx + 1}, col {c_idx + 1} ===")
                    # Show surrounding cell context (same row)
                    row_cells = [c.text[:80] for c in row.cells]
                    for ci, ct in enumerate(row_cells):
                        marker = ">>>" if ci == c_idx else "   "
                        print(f"{marker} [{ci + 1}] {ct}")

    if total_matches == 0:
        print(f"No matches found for: {query}")
    else:
        print(f"\n--- {total_matches} match(es) found ---")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Search within DOCX text")
    parser.add_argument("path", help="Path to DOCX file")
    parser.add_argument("query", help="Search query (regex or literal)")
    parser.add_argument("--context", type=int, default=3,
                        help="Context lines around match (default: 3)")
    parser.add_argument("--literal", action="store_true",
                        help="Treat query as literal string")
    args = parser.parse_args()

    search(args.path, args.query, args.context, args.literal)
